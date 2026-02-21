import os
import re

try:
    import win32com.client as win32
except Exception:  # pragma: no cover - optional dependency
    win32 = None
from docx import Document
from docxtpl import DocxTemplate, RichText

from app.doc_helpers import muat_gambar


class ReportService:
    def __init__(self, templates_dir):
        self.templates_dir = templates_dir

    @staticmethod
    def _normalize_langkah_list(raw_value):
        """
        LOGIKA BARU (ADAPTASI DARI USER):
        Mengubah input menjadi List of Strings dengan format "1.\tLangkah Kerja".
        Menambahkan '##HAPUS##' di awal untuk trik kerapian.
        """
        if not raw_value:
            return ["##HAPUS##", "Langkah kerja tidak diisi."]

        # Ambil list baris teks
        if isinstance(raw_value, list):
            # Jika input dari UI sudah list, ambil teksnya saja jika itu dict
            lines = []
            for item in raw_value:
                if isinstance(item, dict):
                    lines.append(item.get("langkah_kerja", ""))
                else:
                    lines.append(str(item))
        else:
            # Jika input string (dari textarea), pecah berdasarkan Enter
            lines = raw_value.splitlines()

        normalized = []
        nomor_urut = 1

        for line in lines:
            if not isinstance(line, str): continue
            
            # 1. Bersihkan spasi kiri kanan
            teks = line.strip()
            
            # 2. FILTER KERAS (VACUUM CLEANER) a la generator_laporan.py
            # Hapus angka/bullet bawaan user supaya kita bisa nomori ulang
            teks_bersih = teks.lstrip(' *-•.1234567890').strip()
            
            if teks_bersih:
                # 3. FORMATTING RAPI: Angka + Titik + TAB + Teks
                item_baru = f"{nomor_urut}.\t{teks_bersih}"
                normalized.append(item_baru)
                nomor_urut += 1

        # Tambahkan pemicu ##HAPUS## di paling atas
        if not normalized:
            return ["##HAPUS##", "Langkah kerja kosong."]
            
        return ["##HAPUS##"] + normalized
    
    @staticmethod
    def _hapus_baris_hantu(output_path):
        """Menghapus baris yang berisi ##HAPUS##"""
        try:
            doc = Document(output_path)
            for p in list(doc.paragraphs):
                if "##HAPUS##" in p.text:
                    p._element.getparent().remove(p._element)
            doc.save(output_path)
        except Exception:
            pass

    @staticmethod
    def _hapus_paragraf_kosong_spesifik(output_path):
        """
        Fungsi 'Sniper': Hanya menghapus baris kosong di area Langkah Kerja,
        tanpa merusak Cover.
        """
        print("🎯 Membersihkan spasi renggang...")
        try:
            doc = Document(output_path)
            area_aman = False
            paragraphs_to_delete = []
            
            for p in doc.paragraphs:
                teks = p.text.strip()
                # Trigger aman setelah melewati Daftar Isi atau Bab 1
                if "DAFTAR ISI" in teks or "BAB I" in teks or "HASIL PRAKTIKUM" in teks:
                    area_aman = True
                
                if area_aman:
                    # Cek kosong & bukan gambar
                    kosong = not teks
                    xml = p._element.xml
                    ada_gambar = "w:drawing" in xml or "w:object" in xml or "w:inline" in xml
                    
                    if kosong and not ada_gambar:
                        paragraphs_to_delete.append(p)

            for p in paragraphs_to_delete:
                try:
                    p._element.getparent().remove(p._element)
                except:
                    pass
            doc.save(output_path)
        except Exception as e:
            print(f"⚠️ Warning pembersihan: {e}")

    def resolve_template(self, template_choice):
        name = "format-1.docx" if template_choice == "1" else "format-2.docx"
        template_path = os.path.join(self.templates_dir, name)
        if not os.path.exists(template_path):
            raise FileNotFoundError(template_path)
        return template_path

    def build_bab1_context(self, doc, bab1_items, template_choice="1"):
        daftar_sub_bab = []
        counter_gbr_bab1 = 1
        for item in bab1_items:
            label_a = item.get("label_point_a") or (
                "Source Code" if item.get("tipe") == "1" else "Langkah Kerja"
            )
            langkah_list = self._normalize_langkah_list(
                item.get("langkah_list") or item.get("isi_a", "")
            )
            if item.get("tipe") == "2":
                isi_a = langkah_list
            else:
                isi_a = []

            kode_items = item.get("list_kode") or item.get("kode_files", [])
            list_kode_final = []

            total_files = len(kode_items) # Hitung total file dulu

            if total_files > 0:
                for i, d in enumerate(kode_items, 1):
                    # Ambil nama file, prioritas 'judul' lalu 'nama'
                    nama_tampil = d.get("judul") or d.get("nama", "")

                    if total_files > 1:
                        # --- KASUS MULTIPLE CODE ---
                        # Pakai Nomor (1. NamaFile)
                        # Tambah "\n" (Enter) jika ini file ke-2 atau lebih.
                        # Tujuannya: Memutus rantai tabel agar tidak menyatu di Word.
                        prefix = "\n" if i > 1 else ""
                        judul_tampil = f"{prefix}{i}. {nama_tampil}"
                    else:
                        # --- KASUS SINGLE CODE ---
                        # Hapus penanda (Judul jadi string kosong)
                        judul_tampil = ""

                    list_kode_final.append(
                        {"judul": judul_tampil, "isi": d.get("isi", "")}
                    )

            list_gbr = []
            for g in item.get("list_gambar") or item.get("gambar_paths", []):
                obj = muat_gambar(doc, g.get("path", ""))
                if not obj:
                    continue
                list_gbr.append(
                    {
                        "objek_gambar": obj,
                        "caption": g.get("caption", ""),
                        "nomor_tampil": counter_gbr_bab1,
                    }
                )
                counter_gbr_bab1 += 1

            # --- LOGIKA ANALISA (SISTEM LOOPING LIST BARU - BOLEH DIUBAH) ---
            raw_analisa = item.get("isi_analisa") or item.get("analisa", "")
            list_paragraf_analisa = []
            
            if raw_analisa:
                # Pecah berdasarkan baris, bersihkan spasi liar
                paragraphs = [p.strip() for p in raw_analisa.splitlines() if p.strip()]
                for p in paragraphs:
                    # Bersihkan simbol markdown agar teks murni
                    clean_p = p.replace("**", "").replace("*", "").replace("`", "")
                    prefix = "\t" if template_choice == "1" else ""
                    list_paragraf_analisa.append({"teks": prefix + clean_p})
            else:
                prefix = "\t" if template_choice == "1" else ""
                list_paragraf_analisa = [{"teks": prefix + "Analisa belum diisi."}]

            # Masukkan ke dalam daftar_sub_bab
            daftar_sub_bab.append(
                {
                    "judul_sub_bab": item.get("judul_sub_bab", ""),
                    "label_point_a": label_a,
                    "isi_point_a": isi_a,
                    "langkah_list": langkah_list,
                    "list_gambar": list_gbr,
                    # Gunakan variabel baru ini di template Word Anda
                    "list_paragraf_analisa": list_paragraf_analisa, 
                    "list_kode": list_kode_final,
                    # Tetap sertakan isi_analisa original jika masih dibutuhkan di bagian lain
                    "isi_analisa": raw_analisa, 
                }
            )

        return daftar_sub_bab

    @staticmethod
    def _update_toc_word(nama_file):
        print("      ⏳ Sedang melakukan Auto-Update Daftar Isi & Gambar...")
        if win32 is None:
            print("      ⚠️ Win32COM tidak tersedia. Lewati update TOC otomatis.")
            return
        try:
            path_lengkap = os.path.abspath(nama_file)
            word_app = win32.Dispatch("Word.Application")
            word_app.Visible = False
            word_app.DisplayAlerts = False

            doc = word_app.Documents.Open(path_lengkap)
            for toc in doc.TablesOfContents:
                toc.Update()
            doc.Fields.Update()
            doc.Save()
            doc.Close()
            word_app.Quit()
            print("      ✅ Auto-Update Selesai! Dokumen siap cetak.")
        except Exception as e:
            print(
                "      ⚠️ Gagal Auto-Update (Buka file manual lalu tekan Ctrl+A -> F9): "
                f"{e}"
            )

    def build_bab2_context(self, doc, bab2_items, template_choice="1"):
        daftar_tugas = []
        counter_gbr_bab2 = 1

        for item in bab2_items:
            tipe_konten = item.get("tipe_konten") or item.get("tipe", "1")
            judul_tugas = item.get("judul_tugas") or item.get("judul_sub_bab", "")
            isi_deskripsi = item.get("isi_deskripsi") or item.get("isi_point_a") or item.get("isi_a", "")
            
            # --- LOGIKA ANALISA (SISTEM LOOPING LIST BARU - BOLEH DIUBAH) ---
            raw_analisa_tugas = item.get("isi_analisa_tugas") or item.get("isi_analisa") or item.get("analisa", "")
            list_paragraf_analisa_tugas = []
            isi_analisa_final = ""

            if raw_analisa_tugas:
                # Pecah berdasarkan baris, bersihkan spasi liar
                paragraphs = [p.strip() for p in raw_analisa_tugas.splitlines() if p.strip()]
                for p in paragraphs:
                    # Bersihkan simbol markdown agar teks murni
                    clean_p = p.replace("**", "").replace("*", "").replace("`", "")
                    prefix = "\t" if template_choice == "1" else ""
                    list_paragraf_analisa_tugas.append({"teks": prefix + clean_p})
            else:
                prefix = "\t" if template_choice == "1" else ""
                list_paragraf_analisa_tugas = [{"teks": prefix + "Analisa belum diisi."}]

            gambar_items_raw = item.get("gambar_items") or item.get("list_gambar") or item.get("gambar_paths", [])
            list_gbr_tgs = []
            for gambar in gambar_items_raw:
                path_gambar = gambar.get("path", "")
                obj = muat_gambar(doc, path_gambar)
                if not obj:
                    continue
                caption_gambar = gambar.get("caption_gambar") or gambar.get("caption", "")
                list_gbr_tgs.append(
                    {
                        "objek_gambar": obj,
                        "caption": caption_gambar,
                        "nomor_tampil": counter_gbr_bab2,
                    }
                )
                counter_gbr_bab2 += 1

            kode_items_raw = item.get("kode_items") or item.get("list_kode") or item.get("kode_files", [])
            list_kode_final = []
            
            if tipe_konten == "1" and kode_items_raw:
                total_kode = len(kode_items_raw)
                for idx, kode in enumerate(kode_items_raw, 1):
                    # Ambil nama file dengan prioritas yang konsisten
                    nama_file = kode.get("judul_kode") or kode.get("judul") or kode.get("nama_file") or kode.get("nama", "")
                    isi_kode = kode.get("isi_kode") or kode.get("isi", "")

                    if total_kode > 1:
                        # Kasus Multiple: Gunakan penomoran dan prefix Enter untuk memutus tabel
                        prefix = "\n" if idx > 1 else ""
                        judul_tampil = f"{prefix}{idx}. {nama_file}"
                    else:
                        # Kasus Single: Gunakan trik ##HAPUS## agar judul tidak muncul (sama seperti Bab 1)
                        judul_tampil = "##HAPUS##"

                    list_kode_final.append({
                        "judul": judul_tampil, 
                        "isi": isi_kode
                    })

            qa_items = item.get("qa_items") or item.get("qa_list", [])
            qa_questions = []
            qa_answers = []
            qa_list_normalized = []
            for idx, qa in enumerate(qa_items, 1):
                pertanyaan = (qa.get("pertanyaan") or qa.get("q") or "").strip()
                jawaban = (qa.get("jawaban") or qa.get("a") or "").strip()
                if pertanyaan:
                    qa_questions.append(f"{idx}. {pertanyaan}")
                if jawaban:
                    jawaban = jawaban.replace("**", "").replace("*", "")
                    jawaban = jawaban.replace("``", "").replace("`", "")
                    jawaban = jawaban.replace('"', "")

                    qa_answers.append(f"{idx}. {jawaban}")
                if pertanyaan or jawaban:
                    qa_list_normalized.append({"q": pertanyaan, "a": jawaban})

            raw_langkah = item.get("langkah_kerja_items") or item.get("langkah_list") or isi_deskripsi
            langkah_list = self._normalize_langkah_list(raw_langkah)

            if tipe_konten == "3":
                isi_soal = "\n".join(qa_questions)
                isi_jawaban = "\n".join(qa_answers)
                label_point_a = "Pertanyaan & Jawaban"
                isi_point_a = "\n".join(qa_questions)
                isi_analisa_final = ""
            elif tipe_konten == "2":
                isi_soal = isi_deskripsi
                isi_jawaban = raw_analisa_tugas
                label_point_a = "Langkah Kerja"
                isi_point_a = isi_deskripsi
                isi_analisa = raw_analisa_tugas
            else:
                isi_soal = ""
                isi_jawaban = raw_analisa_tugas
                label_point_a = "Source Code"
                isi_point_a = ""
                isi_analisa_final = raw_analisa_tugas

            daftar_tugas.append(
                {
                    "judul_tugas": judul_tugas,
                    "tipe_konten": tipe_konten,
                    "isi_deskripsi": isi_deskripsi,
                    "list_paragraf_analisa_tugas": list_paragraf_analisa_tugas,
                    "kode_items": list_kode_final,
                    "gambar_items": list_gbr_tgs,
                    "qa_items": qa_list_normalized,
                    "langkah_kerja_items": langkah_list,
                    "judul_sub_bab": judul_tugas,
                    "tipe": tipe_konten,
                    "label_point_a": label_point_a,
                    "isi_point_a": isi_point_a,
                    "list_kode": list_kode_final,
                    "isi_analisa": isi_analisa_final,
                    "langkah_list": langkah_list,
                    "qa_list": qa_list_normalized,
                    "isi_soal": isi_soal,
                    "isi_jawaban": isi_jawaban,
                    "ada_gambar": len(list_gbr_tgs) > 0,
                    "list_gambar": list_gbr_tgs,
                }
            )

        return daftar_tugas
    
    @staticmethod
    def _brute_force_delete_empty_rows(output_path):
        """
        Fungsi 'Preman': Membuka file hasil generate, masuk ke setiap tabel,
        dan menghapus paksa baris (row) yang tidak ada teksnya sama sekali.
        """
        print(f"🕵️  Memulai operasi pembersihan baris kosong di {output_path}...")
        try:
            # Buka ulang dokumen yang baru saja disimpan
            doc_fix = Document(output_path)
            
            total_hapus = 0
            
            # Loop semua tabel yang ada di dokumen
            for table in doc_fix.tables:
                # Kita looping DARI BAWAH ke ATAS (reversed).
                # Kenapa? Karena kalau hapus dari atas, indeks barisnya bakal geser dan bikin error.
                for row in reversed(table.rows):
                    # Gabungkan teks dari semua sel di baris itu
                    row_text = ""
                    for cell in row.cells:
                        row_text += cell.text.strip()
                    
                    # Cek: Kalau gabungan teksnya kosong melompong (cuma spasi/enter doang)
                    if not row_text:
                        # HAPUS BARISNYA!
                        row._element.getparent().remove(row._element)
                        total_hapus += 1

            # Simpan kembali (Overwrite)
            doc_fix.save(output_path)
            print(f"✅ Selesai! Berhasil menghapus {total_hapus} baris hantu yang renggang.")
            
        except Exception as e:
            print(f"⚠️ Gagal melakukan pembersihan: {e}")

    def render_report(
        self, template_choice, cover, bab1_items, bab2_items, kesimpulan, output_path
    ):
        template_path = self.resolve_template(template_choice)
        doc = DocxTemplate(template_path)

        if kesimpulan:
            list_paragraf_kesimpulan = []
            for p in kesimpulan.splitlines():
                clean_p = p.strip()
                if clean_p:
                    prefix = "\t" if template_choice == "1" else ""
                    list_paragraf_kesimpulan.append({"teks": prefix + clean_p})
        else:
            list_paragraf_kesimpulan = [{"teks": "Kesimpulan belum diisi."}]

        daftar_sub_bab = self.build_bab1_context(doc, bab1_items, template_choice)
        daftar_tugas = self.build_bab2_context(doc, bab2_items, template_choice)

        context = {
            **cover,
            "daftar_sub_bab": daftar_sub_bab,
            "daftar_tugas": daftar_tugas,
            "list_kesimpulan": list_paragraf_kesimpulan,
        }

        doc.render(context)
        doc.save(output_path)

        self._hapus_baris_hantu(output_path)
        self._hapus_paragraf_kosong_spesifik(output_path)
        self._brute_force_delete_empty_rows(output_path)
        self._update_toc_word(output_path)
